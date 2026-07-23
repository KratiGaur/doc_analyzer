from io import BytesIO
from pathlib import Path


def _load_pdf_reader():
    try:
        from PyPDF2 import PdfReader
    except ImportError as exc:
        raise RuntimeError("PDF support is not installed. Install PyPDF2 to read PDFs.") from exc
    return PdfReader


def _load_document_module():
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX support is not installed. Install python-docx to read DOCX files.") from exc
    return Document


def _load_pil():
    try:
        from PIL import Image, UnidentifiedImageError
    except ImportError as exc:
        raise RuntimeError("Image OCR support is not installed. Install Pillow to read images.") from exc
    return Image, UnidentifiedImageError


def _load_tesseract():
    try:
        import pytesseract
    except ImportError as exc:
        raise RuntimeError("OCR support is not installed. Install pytesseract and Tesseract to read scans.") from exc
    return pytesseract


def _extract_pdf_image_bytes(page) -> list[bytes]:
    """Best-effort extraction of embedded image bytes from a PDF page."""
    image_bytes_list: list[bytes] = []

    page_images = getattr(page, "images", None)
    if page_images:
        for image in page_images:
            data = getattr(image, "data", None) or getattr(image, "_data", None)
            if data:
                image_bytes_list.append(data)
        if image_bytes_list:
            return image_bytes_list

    try:
        resources = page.get("/Resources")
        if not resources:
            return image_bytes_list

        xobject = resources.get("/XObject")
        if not xobject:
            return image_bytes_list

        xobject = xobject.get_object()
        for obj in xobject.values():
            try:
                resolved = obj.get_object()
                if resolved.get("/Subtype") == "/Image":
                    data = resolved.get_data()
                    if data:
                        image_bytes_list.append(data)
            except Exception:
                continue
    except Exception:
        return image_bytes_list

    return image_bytes_list


def _ocr_image_bytes(image_bytes: bytes) -> str:
    Image, _ = _load_pil()
    pytesseract = _load_tesseract()

    image = Image.open(BytesIO(image_bytes))
    image = image.convert("RGB")
    return pytesseract.image_to_string(image).strip()


def extract_text_from_pdf(pdf_path: str) -> str:
    """Extract text from a PDF. Fall back to OCR for scanned/image-only pages."""
    PdfReader = _load_pdf_reader()
    reader = PdfReader(pdf_path)

    extracted_text = []
    ocr_text = []

    for page in reader.pages:
        page_text = (page.extract_text() or "").strip()
        if page_text:
            extracted_text.append(page_text)
            continue

        for image_bytes in _extract_pdf_image_bytes(page):
            try:
                image_text = _ocr_image_bytes(image_bytes)
            except RuntimeError as exc:
                raise RuntimeError(
                    "This PDF looks scanned, but OCR dependencies are missing. "
                    "Install Pillow, pytesseract, and Tesseract to read image-based PDFs."
                ) from exc
            if image_text:
                ocr_text.append(image_text)

    combined_text = "\n".join(extracted_text).strip()
    if combined_text:
        return combined_text

    ocr_combined = "\n".join(ocr_text).strip()
    if ocr_combined:
        return ocr_combined

    raise RuntimeError(
        "No extractable text was found in this PDF. "
        "If it is a scanned document, install OCR support (Pillow, pytesseract, and Tesseract)."
    )


def extract_text_from_docx(docx_path: str) -> str:
    """Extract text from DOCX files."""
    try:
        Document = _load_document_module()
        document = Document(docx_path)
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

        table_text = []
        for table in document.tables:
            for row in table.rows:
                table_text.extend([cell.text.strip() for cell in row.cells if cell.text.strip()])

        return "\n".join(paragraphs + table_text).strip()
    except Exception as exc:
        raise RuntimeError(f"Unable to parse DOCX file: {exc}") from exc


def extract_text_from_txt(txt_path: str) -> str:
    """Read plain text files."""
    try:
        with open(txt_path, "r", encoding="utf-8", errors="replace") as txt_file:
            return txt_file.read().strip()
    except Exception as exc:
        raise RuntimeError(f"Unable to read TXT file: {exc}") from exc


def extract_text_from_image(image_path: str) -> str:
    """Extract text from an image file using OCR."""
    try:
        Image, UnidentifiedImageError = _load_pil()
        pytesseract = _load_tesseract()
        image = Image.open(image_path)
        image = image.convert("RGB")
        return pytesseract.image_to_string(image).strip()
    except UnidentifiedImageError as exc:
        raise RuntimeError("The uploaded image file could not be read.") from exc
    except RuntimeError:
        raise
    except Exception as exc:
        raise RuntimeError(f"Unable to parse image file: {exc}") from exc


def extract_text_from_file(file_path: str) -> str:
    """Choose the correct extractor based on file extension."""
    extension = Path(file_path).suffix.lower()

    if extension == ".pdf":
        return extract_text_from_pdf(file_path)
    if extension == ".docx":
        return extract_text_from_docx(file_path)
    if extension == ".txt":
        return extract_text_from_txt(file_path)
    if extension in {".png", ".jpg", ".jpeg"}:
        return extract_text_from_image(file_path)

    raise ValueError(f"Unsupported file format: {extension}")
